"""
Generates OnChainCarFacts-Anchor-Review.docx — a short review doc covering:
  1. The three anchor contract shape options
  2. A side-by-side L1 (Ethereum) vs L2 (Base) cost comparison

Re-run this script after editing to regenerate the .docx. All figures are
ballpark and uncited — flagged [S?] per the project's attribution convention.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "OnChainCarFacts-Anchor-Review.docx"

doc = Document()

# Base style tweaks
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---------- Title ----------
title = doc.add_heading("OnChainCarFacts — Anchor Design Review", level=0)
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("Problem 1 (Record Anchoring) — for team review, 2026-04-15")
subtitle_run.italic = True
subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ---------- Purpose ----------
doc.add_heading("Purpose", level=1)
doc.add_paragraph(
    "This document presents two decisions for team review: (1) the shape of the "
    "on-chain anchoring contract, and (2) where to deploy it (Ethereum L1 vs. "
    "Base L2) and how to show the cost difference as evidence in the final "
    "write-up. It is a narrow follow-up to the Problem 1 proposal "
    "(docs/problem1-record-anchoring.md) and does not revisit out-of-scope "
    "topics (attestation, reputation, privacy, fraud detection)."
)

doc.add_paragraph(
    "Attribution tags used below follow the convention in docs/SOURCES.md: "
    "[S#] = cited source, [S?:tag] = source needed before final submission, "
    "[CC] = framing co-developed with Claude, [OWN] = our own analysis. "
    "All dollar figures in this document are ballpark and must be replaced "
    "with measured values (and cited gas-price / ETH-price sources) before "
    "they appear in the final report."
)

# ---------- Contract shape ----------
doc.add_heading("1. Contract shape — three options", level=1)

doc.add_paragraph(
    "All three options put a cryptographic commitment to a record on chain and "
    "let a later verifier prove the record existed by the anchoring block's "
    "timestamp. They differ in what is stored on chain, how gas cost scales "
    "with record volume, and how much Merkle-tree machinery lives off chain. "
    "[OWN]"
)

headers = [
    "Dimension",
    "Option 1: anchor(bytes32)",
    "Option 2: anchorBatch(bytes32[])",
    "Option 3: anchorRoot(bytes32)",
]

rows = [
    [
        "What is on chain",
        "One hash per record; one transaction per record.",
        "N hashes per transaction (each emitted as its own event).",
        "One Merkle root per transaction, covering N records. Leaves live off chain.",
    ],
    [
        "Gas amortization",
        "None. Every record pays the full base transaction cost.",
        "Base transaction cost (~21k gas) is shared across N records. Per-hash storage / event cost still paid.",
        "A single 32-byte commitment on chain regardless of N. Cheapest per record at scale. [S?:merkle-tree]",
    ],
    [
        "Verifier workflow",
        "Look up hash by ID, compare to the record's recomputed hash.",
        "Same as Option 1.",
        "Recompute leaf hash, walk a Merkle proof (log2(N) sibling hashes) up to the stored root. [S?:merkle-tree]",
    ],
    [
        "Off-chain complexity",
        "Trivial.",
        "Trivial.",
        "Must build and retain the Merkle tree; must be able to serve proofs at verify time.",
    ],
    [
        "Failure mode if off-chain data is lost",
        "Any record you still possess can still be verified against its on-chain hash.",
        "Same as Option 1.",
        "Root becomes unverifiable: without the leaves you cannot produce proofs. Must be discussed in the threat model. [OWN]",
    ],
    [
        "Best fit",
        "Small demos or very low record volume.",
        "Medium volume where simple per-record lookups are still desired.",
        "Production-scale anchoring; the academically interesting test of what batching actually buys.",
    ],
]

table = doc.add_table(rows=1 + len(rows), cols=len(headers))
table.style = "Light Grid Accent 1"
table.autofit = True

hdr = table.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True

for ri, row in enumerate(rows, start=1):
    cells = table.rows[ri].cells
    for ci, val in enumerate(row):
        cells[ci].text = val

doc.add_heading("Recommendation", level=2)
doc.add_paragraph(
    "Build Option 1 and Option 3 side by side; skip Option 2. [OWN]"
)
doc.add_paragraph(
    "Option 1 gives us the clean, easy-to-explain demo path: submit a record, "
    "see its hash land on chain, verify it later. Option 3 gives us the "
    "interesting scaling story: a single on-chain transaction commits to many "
    "records, and the verifier does the Merkle math. Option 2 sits in the "
    "middle and does not teach anything the other two do not. [OWN]"
)

# ---------- Cost comparison ----------
doc.add_heading("2. Cost comparison — Ethereum L1 vs. Base L2", level=1)

doc.add_paragraph(
    "The table below gives ballpark deployment and per-anchor costs on "
    "Ethereum mainnet vs. Base mainnet. Assumptions (all still need citation): "
    "Ethereum gas price ~15 gwei, ETH price ~$3,000, Base L2 cost ~1000–2000x "
    "cheaper per operation than Ethereum L1. Testnet costs are zero on both "
    "chains (faucet ETH). [S?:base-l2] [S?:l2-rollup-concept]"
)

cost_headers = [
    "Scenario",
    "Ethereum L1 (Sepolia / mainnet)",
    "Base L2 (Base Sepolia / mainnet)",
]

cost_rows = [
    [
        "Bare transaction cost (21,000 gas)",
        "~$0.95",
        "fractions of a cent",
    ],
    [
        "Single-hash anchor (Option 1; ~43,000 gas)",
        "~$1.50 – $3.00",
        "~$0.001 – $0.003",
    ],
    [
        "Batch of 100 hashes via anchorBatch (Option 2)",
        "~$50 – $100",
        "low cents",
    ],
    [
        "Batch of 100 hashes via anchorRoot (Option 3)",
        "~$1.50 – $3.00 (only the root is stored)",
        "~$0.001 – $0.003",
    ],
    [
        "Testnet cost (Sepolia / Base Sepolia)",
        "$0 (faucet ETH)",
        "$0 (faucet ETH)",
    ],
]

ct = doc.add_table(rows=1 + len(cost_rows), cols=len(cost_headers))
ct.style = "Light Grid Accent 1"
ct.autofit = True

hdr = ct.rows[0].cells
for i, h in enumerate(cost_headers):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True

for ri, row in enumerate(cost_rows, start=1):
    cells = ct.rows[ri].cells
    for ci, val in enumerate(row):
        cells[ci].text = val

doc.add_heading("Observations", level=2)

doc.add_paragraph(
    "On Ethereum L1, Option 3 (Merkle root) is the only economically viable "
    "design at any realistic record volume. Per-record anchoring or per-record "
    "batched events cost dollars each; a protocol that expects thousands of "
    "records per day cannot absorb that cost. [OWN]",
    style="List Bullet",
)
doc.add_paragraph(
    "On Base L2, Options 1 and 2 become cheap enough that the extra off-chain "
    "complexity of Option 3 is not obviously worth it for the property we are "
    "testing. The argument for L2 is therefore not just \"lower gas\" — it is "
    "\"lower gas, which unlocks a simpler design.\" That second-order effect is "
    "the more interesting finding for the write-up. [OWN] [CC] (framing "
    "suggested by Claude; accepted after review.)",
    style="List Bullet",
)
doc.add_paragraph(
    "Deploying the same contract on both Ethereum Sepolia and Base Sepolia "
    "and anchoring the same sample batch on each lets us replace every "
    "ballpark figure above with measured gas, then multiply by current "
    "mainnet gas-price / ETH-price data (with citations) for the final "
    "write-up. [OWN]",
    style="List Bullet",
)

# ---------- Decisions to ratify ----------
doc.add_heading("3. Decisions to ratify", level=1)
doc.add_paragraph(
    "Proposed direction, pending team agreement. [OWN]"
)
doc.add_paragraph(
    "Contract shape: implement Option 1 and Option 3; skip Option 2.",
    style="List Number",
)
doc.add_paragraph(
    "Deployments: same contract deployed to Ethereum Sepolia and Base Sepolia, "
    "so the L1-vs-L2 cost comparison is measured on our own transactions, not "
    "estimated from secondary data.",
    style="List Number",
)
doc.add_paragraph(
    "Phase 0 (the existing NHTSA VIN-lookup demo) is left as-is. The anchoring "
    "work lives in a separate minimal app so the two concerns stay cleanly "
    "separated in the final submission.",
    style="List Number",
)

# ---------- Attribution footer ----------
doc.add_heading("Attribution", level=1)
doc.add_paragraph(
    "Every external or non-baseline claim in this document is tagged. Any [S?] "
    "marker is a placeholder and must be replaced with a real [S#] entry in "
    "docs/SOURCES.md before final submission. Class-covered baseline material "
    "(smart contract fundamentals, on-chain storage) is not tagged. Material "
    "beyond the class baseline — L2 rollups, specific testnets, Merkle trees, "
    "certificate-transparency-style logs — is tagged and must be cited."
)

doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
