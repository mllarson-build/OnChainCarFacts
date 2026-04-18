"""
Generates OnChainCarFacts-Team-Review-Guide.docx from the content of
docs/team-review-guide.md. The Markdown remains the source of truth;
re-run this script after editing the .md to regenerate the .docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor

OUTPUT = "OnChainCarFacts-Team-Review-Guide.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

MONO = "Consolas"


def mono(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(10)
    return run


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def table(headers, rows, style_name="Light Grid Accent 1"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = style_name
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return t


# ---------- Title ----------
doc.add_heading("Team Review Guide — Record Anchoring (Problem 1)", level=0)
sub = doc.add_paragraph()
run = sub.add_run(
    "Audience: teammates and the instructor who want to (a) review the code, "
    "(b) run the tests, and (c) understand why the design works."
)
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph(
    "This guide is deliberately short. For milestone-by-milestone scope see "
    "docs/problem1-build-plan.md. For the design trade-offs see "
    "docs/problem1-solution-tradeoffs.md. Attribution tags used below follow "
    "docs/SOURCES.md."
)

# ---------- 1. What is deployed ----------
doc.add_heading("1. What is deployed right now", level=1)

table(
    ["Item", "Value"],
    [
        ["Contract", "RecordAnchor.sol"],
        ["Network", "Base Sepolia (chain id 84532)"],
        ["Contract address", "0x1380b4cBC2Bdb10e46F720124C8174bc363aE4bF"],
        ["Deploy tx",
         "0xf566c99da45feee5da821b55547004eb25b3d92e5932c0c3ab076961cb4fbfcd"],
        ["Sample anchor tx",
         "0xad6fd03965116b763ad9e7e34a86d3895de7a5a8dae5a5997ee37c88c4a02676"],
        ["Anchor gas used", "90,594"],
    ],
)

doc.add_paragraph("Explorer (clickable):")
doc.add_paragraph(
    "Contract — https://sepolia.basescan.org/address/"
    "0x1380b4cBC2Bdb10e46F720124C8174bc363aE4bF",
    style="List Bullet",
)
doc.add_paragraph(
    "Anchor tx — https://sepolia.basescan.org/tx/"
    "0xad6fd03965116b763ad9e7e34a86d3895de7a5a8dae5a5997ee37c88c4a02676",
    style="List Bullet",
)

doc.add_paragraph(
    "You can verify the anchor without any of our code by opening the "
    "contract on Basescan, clicking “Read Contract” → getAnchor, and "
    "pasting the record hash. It will return the timestamp, block, and "
    "submitter. That is the core claim of the project: a third party can "
    "independently verify a record existed at a given time, using only a "
    "public block explorer."
)

# ---------- 2. How to review ----------
doc.add_heading("2. How to review the code (5-minute path)", level=1)
doc.add_paragraph(
    "Read these files in order. The whole pipeline is under 300 lines."
)
review_steps = [
    ("anchor/src/types.ts", "the record schema (v1.0). What we hash."),
    ("anchor/src/canonicalize.ts", "deterministic byte encoding of a record."),
    ("anchor/src/hash.ts", "keccak256(canonicalize(record))."),
    ("anchor/contracts/RecordAnchor.sol", "the contract (~110 lines)."),
    ("anchor/test/",
     "the tests, especially the hash-parity test that proves our "
     "TypeScript hash matches Solidity’s."),
    ("anchor/scripts/deploy.ts and anchor/scripts/anchor-one.ts",
     "how we deployed and how we wrote one record to Base Sepolia."),
]
for i, (path, desc) in enumerate(review_steps, 1):
    p = doc.add_paragraph(style="List Number")
    mono(p, path)
    p.add_run(f" — {desc}")

# ---------- 3. Run locally ----------
doc.add_heading("3. How to run it locally", level=1)
doc.add_paragraph(
    "Prereq: Node 20+, git, a terminal. No blockchain knowledge required "
    "for the local run — it uses an in-memory Hardhat node."
)
code_block(
    "git clone <this repo>\n"
    "cd onchaincarfacts/anchor\n"
    "npm install\n"
    "npx hardhat compile\n"
    "npx hardhat test"
)
doc.add_paragraph(
    "Expected: all tests pass, including hash-parity (JS and Solidity "
    "produce the same bytes32) and Merkle-root verification."
)

doc.add_heading("Optional: deploy to a public testnet yourself", level=2)
doc.add_paragraph(
    "You do not need to do this to grade the project — the deployment "
    "above is already public and immutable. If you want to repeat it:"
)
p = doc.add_paragraph(style="List Number")
mono(p, "cp .env.example .env")
p.add_run(" and fill in:")
doc.add_paragraph(
    "BASE_SEPOLIA_RPC_URL — e.g. https://sepolia.base.org (public) or an "
    "Alchemy/Infura endpoint.",
    style="List Bullet 2",
)
doc.add_paragraph(
    "DEPLOYER_PRIVATE_KEY — a testnet-only key. Fund it from a Base "
    "Sepolia faucet (~0.01 ETH is plenty).",
    style="List Bullet 2",
)
p = doc.add_paragraph(style="List Number")
mono(p, "npx hardhat run scripts/deploy.ts --network baseSepolia")
p = doc.add_paragraph(style="List Number")
mono(p, "npx hardhat run scripts/anchor-one.ts --network baseSepolia")
doc.add_paragraph(
    "Your own deployment will land at a different address; "
    "deployments.json records both the shared deployment and anything "
    "you add locally."
)

# ---------- 4. How it works ----------
doc.add_heading("4. How it works (the pipeline)", level=1)
code_block(
    " vehicle record (JSON)\n"
    "        │\n"
    "        ▼  canonicalize.ts — encode fields in a fixed order\n"
    " canonical bytes\n"
    "        │\n"
    "        ▼  hash.ts — keccak256\n"
    " 32-byte hash\n"
    "        │\n"
    "        ▼  contract.anchor(hash)   (sent as an Ethereum tx)\n"
    " on-chain mapping: hash → (timestamp, blockNumber, submitter)\n"
    "        │\n"
    "        ▼  anyone later:  contract.getAnchor(hash) → same tuple\n"
    " independently verifiable proof of existence"
)

doc.add_paragraph("Two modes are implemented in the same contract:")
doc.add_paragraph(
    "Option 1 — single-hash anchor (anchor / getAnchor). One tx per "
    "record. Simple, but one tx per record is ~90k gas.",
    style="List Bullet",
)
doc.add_paragraph(
    "Option 3 — Merkle-root anchor (anchorRoot / verifyLeaf). You build "
    "a Merkle tree over many record hashes off-chain, anchor only the "
    "32-byte root, and prove any individual leaf later with a short "
    "proof. One tx covers arbitrarily many records. [OWN — choice to "
    "implement both so M6 can measure the cost delta.]",
    style="List Bullet",
)

doc.add_heading("Why the hash-parity test matters", level=2)
doc.add_paragraph(
    "The contract never sees the raw record — only the hash. If our "
    "off-chain TypeScript hash didn’t match what the contract would "
    "produce from the same record, a verifier couldn’t reconstruct the "
    "hash and the whole system is worthless. The canonicalHash view "
    "function on the contract exists solely to be called from a test "
    "with the same fields the TypeScript code encoded; both must return "
    "the same bytes32. [OWN]"
)

# ---------- 5. Why it works ----------
doc.add_heading("5. Why it works (the cryptography and the chain)", level=1)
doc.add_paragraph(
    "Three properties hold together, and each has a well-understood "
    "source."
)

doc.add_heading("5.1  keccak256 is collision-resistant", level=2)
doc.add_paragraph(
    "A 256-bit cryptographic hash. Finding two different inputs that "
    "produce the same output is computationally infeasible. This is why "
    "a 32-byte hash is a faithful “fingerprint” of the record — change "
    "one character of the VIN or one digit of the mileage and the hash "
    "changes completely. [S?:keccak256 — Ethereum Yellow Paper / "
    "Solidity docs]"
)

doc.add_heading("5.2  Ethereum ordering gives a trusted timestamp", level=2)
doc.add_paragraph(
    "When the contract stores block.timestamp at anchor time, that "
    "timestamp is agreed on by the entire network as part of consensus. "
    "No single party (including us, the submitter) can rewrite it after "
    "the fact without rewriting Ethereum history — which the economic "
    "security of the chain makes impractical. This is why the proof is "
    "“trustless”: the verifier doesn’t have to trust us, only the chain. "
    "[S?:ethereum-consensus]"
)

doc.add_heading("5.3  Base Sepolia inherits Ethereum’s security model", level=2)
doc.add_paragraph(
    "We deploy to Base Sepolia rather than Ethereum mainnet for cost "
    "reasons: gas on an L2 testnet is free (faucet ETH) and real Base "
    "mainnet txs are ~10–100× cheaper than L1. Base is an optimistic "
    "rollup — transactions execute on Base, then batches are posted back "
    "to Ethereum L1, where the data is available for anyone to re-derive "
    "Base’s state. The security argument for production would be “Base "
    "inherits L1 data availability and settlement”; for this project the "
    "testnet is sufficient to demonstrate the pattern. [S?:base-l2] "
    "[S?:l2-rollup-concept]"
)

doc.add_heading("5.4  The pattern has a well-known precedent", level=2)
doc.add_paragraph(
    "What we are doing — hashing records off-chain and anchoring hashes "
    "to a tamper-evident public log — is the same pattern used by "
    "Certificate Transparency, which logs every TLS certificate issued "
    "by a trusted CA so that mis-issuance is publicly detectable. That "
    "system has been running since 2013 and is now required by major "
    "browsers. The blockchain replaces the append-only log server; "
    "everything else is analogous. [S?:certificate-transparency — "
    "RFC 6962]"
)

# ---------- 6. Sources ----------
doc.add_heading("6. Sources to read for deeper understanding", level=1)
doc.add_paragraph(
    "These are reading-list pointers, not yet fully cited entries. "
    "Confirmed citations will land in docs/SOURCES.md as they’re read."
)

table(
    ["Topic", "Suggested source", "Why read it"],
    [
        ["What keccak256 actually computes",
         "Ethereum Yellow Paper, appendix on hashing; Solidity docs",
         "Formal definition of the hash we rely on"],
        ["Why blockchain timestamps hold",
         "Ethereum consensus spec / any intro-to-PoS explainer",
         "Grounds the “trustless timestamp” claim"],
        ["L2 rollups vs. L1",
         "ethereum.org/en/developers/docs/scaling/optimistic-rollups/",
         "Why Base Sepolia is cheap and still credible"],
        ["Base specifically",
         "docs.base.org",
         "Confirms Base is Coinbase’s OP-Stack rollup"],
        ["Certificate Transparency",
         "RFC 6962",
         "The canonical precedent for hash anchoring"],
        ["Merkle trees",
         "Any cryptography textbook; en.wikipedia.org/wiki/Merkle_tree",
         "Background for Option 3 (batched anchoring)"],
        ["Solidity abi.encode determinism",
         "Solidity language docs, ABI section",
         "Why canonicalization + abi.encode is stable"],
    ],
)

doc.add_paragraph(
    "Every [S?] tag in this doc becomes a real [S#] entry in "
    "docs/SOURCES.md as each source is read and verified before final "
    "submission. [OWN]"
)

# ---------- 7. Not in scope ----------
doc.add_heading("7. What’s not in scope (so reviewers don’t ask)", level=1)
doc.add_paragraph(
    "Per the 2026-04-15 scope narrow (docs/problem1-build-plan.md), this "
    "project is only record anchoring. The following are explicitly out "
    "of scope and live in the solution-tradeoffs doc as future work: "
    "attestations / source reputation, privacy-preserving record "
    "storage, fraud detection, and a consumer-facing product. [OWN]"
)

doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
